import os
import io
import tempfile
import logging
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
import win32com.client
import pythoncom
import fitz  # PyMuPDF
from PIL import Image
from difflib import SequenceMatcher
from datetime import datetime
from anthropic import Anthropic
import json
from typing import Dict, Any, Optional, Tuple, Union
import hashlib
from models import JobInfo, CandidateInfo
import docx  # python-docx for DOCX handling
import re

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Anthropic client with error handling
try:
    ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
    if not ANTHROPIC_API_KEY:
        logger.warning("ANTHROPIC_API_KEY not found in environment variables")
        anthropic_client = None
    else:
        anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)
except Exception as e:
    logger.error(f"Failed to initialize Anthropic client: {str(e)}")
    anthropic_client = None

def extract_text_from_pdf(file_bytes):
    try:
        logger.info("Extracting text from PDF")
        is_image_based = False
        text = ""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text()
            if len(text.strip()) < 100:
                logger.info("Limited text found in PDF, trying OCR...")
                is_image_based = True
                images = convert_from_bytes(file_bytes)
                text = ""
                for i, img in enumerate(images):
                    text += pytesseract.image_to_string(img)
                    logger.info(f"OCR processed page {i+1}/{len(images)}")
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {str(e)}, falling back to OCR")
            is_image_based = True
            images = convert_from_bytes(file_bytes)
            text = ""
            for i, img in enumerate(images):
                text += pytesseract.image_to_string(img)
                logger.info(f"OCR processed page {i+1}/{len(images)}")

        logger.info(f"PDF extraction completed: {'image-based' if is_image_based else 'text-based'}")
        return text, is_image_based

    except Exception as e:
        logger.error(f"PDF text extraction failed: {str(e)}")
        raise

def extract_text_from_docx(file_bytes):
    try:
        logger.info("Extracting text from DOCX")
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text, False
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        return "", False

def extract_image_text(image_bytes):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image).strip()
    except Exception as e:
        logger.warning(f"Image OCR error: {str(e)}")
        return ""

def convert_docx_to_pdf(docx_bytes):
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            docx_path = docx_temp.name
            docx_temp.write(docx_bytes)

        pdf_path = docx_path.replace('.docx', '.pdf')
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            return pdf_bytes
        finally:
            word.Quit()
            try:
                os.unlink(docx_path)
                os.unlink(pdf_path)
            except:
                pass
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {str(e)}")
        return None

def remove_duplicates(text: str, threshold: float = 0.92) -> str:
    if not text:
        return ""
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    unique_paragraphs = []
    for p in paragraphs:
        if len(p) < 15:
            unique_paragraphs.append(p)
            continue
        is_duplicate = False
        for existing in unique_paragraphs:
            if len(existing) < 15:
                continue
            if SequenceMatcher(None, p, existing).ratio() > threshold:
                is_duplicate = True
                break
        if not is_duplicate:
            unique_paragraphs.append(p)
    return '\n\n'.join(unique_paragraphs)

def calculate_parse_score(text):
    if not text:
        return 0
    word_count = len(text.split())
    if word_count < 50:
        return max(10, int(word_count / 5))
    paragraphs = text.split('\n\n')
    unique_words = set(word.lower() for word in text.split())
    vocab_richness = len(unique_words) / word_count if word_count else 0
    job_terms = ['experience', 'skills', 'requirements', 'responsibilities', 'qualifications', 
                 'education', 'salary', 'benefits', 'position', 'job', 'work', 'team']
    term_matches = sum(1 for term in job_terms if term.lower() in text.lower())
    term_score = min(40, term_matches * 4)
    structure_score = min(30, len(paragraphs))
    vocab_score = min(20, int(vocab_richness * 100))
    return max(10, min(100, term_score + structure_score + vocab_score))

def clean_text(text: str) -> str:
    """Clean and deduplicate text while maintaining paragraph order."""
    if not text:
        return ""
    
    # Split into paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n")]
    
    # Remove empty paragraphs and deduplicate while maintaining order
    seen = set()
    cleaned_paragraphs = []
    for p in paragraphs:
        if p and p not in seen:
            seen.add(p)
            cleaned_paragraphs.append(p)
    
    return "\n\n".join(cleaned_paragraphs)

from anthropic.types import Message

def format_with_claude(text: str) -> str:
    if not anthropic_client:
        logger.warning("Anthropic not initialized. Skipping formatting.")
        return text

    try:
        logger.info("Sending text to Claude for formatting")
        response: Message = anthropic_client.messages.create(
            model="claude-3-sonnet-20240229",
            temperature=0.2,
            max_tokens=4000,
            system="You are an assistant that formats job descriptions for clarity and structure. Do not rewrite or change meaning.",
            messages=[
                {"role": "user", "content": f"Format this job description cleanly without changing its content:\n\n{text}"}
            ]
        )

        if response and hasattr(response, "content") and isinstance(response.content, list):
            formatted_blocks = [part.text.strip() for part in response.content if part.type == "text"]
            formatted_text = "\n\n".join(formatted_blocks)
            logger.info("Received formatted text from Claude")
            return formatted_text if formatted_text else text

        logger.warning("Claude response was empty or malformed")
        return text

    except Exception as e:
        logger.error(f"Claude formatting failed: {e}")
        return text

def extract_candidate_info(text: str) -> dict:
    """
    Extract structured candidate information from CV text using Claude.
    
    Args:
        text: The extracted text from the CV
        
    Returns:
        dict: Structured candidate information
    """
    if not anthropic_client:
        logger.warning("Anthropic client not available, returning empty candidate info")
        return post_process_extracted_info({})
    
    try:
        prompt = f"""Extract candidate information from the following CV text and return it in the specified JSON schema. 
Pay special attention to work experience sections. Look for job titles, companies, dates, and responsibilities.
If a field is missing from the CV, still include it with null, "", or [] as appropriate. Do not skip any fields.

CV Text:
{text}

Return the information in this exact JSON schema:
{{
    "full_name": "string or null",
    "email": "string or null",
    "phone": "string or null",
    "location": "string or null",
    "linkedin": "string or null",
    "github": "string or null",
    "summary": "string or null",
    "skills": ["string"],
    "education": [
        {{
            "degree": "string",
            "institution": "string",
            "year_completed": "string"
        }}
    ],
    "experience": [
        {{
            "job_title": "string",
            "company": "string",
            "duration": "string (e.g., '2015-2020', '2 years', '6 months')",
            "responsibilities": ["string"]
        }}
    ],
    "certifications": [
        {{
            "name": "string",
            "issuer": "string",
            "year": "string"
        }}
    ],
    "languages": ["string"],
    "availability": "string or null"
}}

Important:
1. For experience, look for sections like 'Work Experience', 'Professional Experience', 'Employment History'
2. Extract dates in any format (e.g., '2015-2020', '2 years', '6 months', 'Jan 2015 - Dec 2020')
3. Include all job positions, even if dates are not clear
4. For responsibilities, extract bullet points or paragraphs describing the role

Return only the JSON object, no other text or explanation."""

        response = anthropic_client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )
        
        if not response.content:
            logger.warning("Empty response from Claude for candidate info extraction")
            return post_process_extracted_info({})
            
        try:
            # Extract the JSON from the response
            content = response.content[0].text
            # Find the first { and last } to extract the JSON
            start = content.find('{')
            end = content.rfind('}') + 1
            if start == -1 or end == 0:
                logger.warning("No JSON found in Claude response")
                return post_process_extracted_info({})
                
            json_str = content[start:end]
            extracted_info = json.loads(json_str)
            
            # Log the extracted experience for debugging
            if extracted_info.get("experience"):
                logger.info(f"Extracted experience: {json.dumps(extracted_info['experience'], indent=2)}")
            else:
                logger.warning("No experience extracted from CV")
                
            return post_process_extracted_info(extracted_info)
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse Claude response as JSON: {e}")
            return post_process_extracted_info({})
            
    except Exception as e:
        logger.error(f"Error extracting candidate info with Claude: {e}")
        return post_process_extracted_info({})

def post_process_extracted_info(info: dict) -> dict:
    """
    Ensure all expected fields exist in the extracted candidate information.
    
    Args:
        info: The extracted candidate information
        
    Returns:
        dict: Post-processed candidate information with all fields
    """
    default_info = {
        "full_name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin": None,
        "github": None,
        "summary": None,
        "skills": [],
        "education": [],
        "experience": [],
        "certifications": [],
        "languages": [],
        "availability": None
    }
    
    # Update with any existing values from info
    for key, value in info.items():
        if key in default_info:
            if isinstance(value, list):
                # Ensure list items have the correct structure
                if key == "education":
                    default_info[key] = [
                        {
                            "degree": item.get("degree", ""),
                            "institution": item.get("institution", ""),
                            "year_completed": item.get("year_completed", "")
                        }
                        for item in value
                    ]
                elif key == "experience":
                    default_info[key] = [
                        {
                            "job_title": item.get("job_title", ""),
                            "company": item.get("company", ""),
                            "duration": item.get("duration", ""),
                            "responsibilities": item.get("responsibilities", [])
                        }
                        for item in value
                    ]
                elif key == "certifications":
                    default_info[key] = [
                        {
                            "name": item.get("name", ""),
                            "issuer": item.get("issuer", ""),
                            "year": item.get("year", "")
                        }
                        for item in value
                    ]
                else:
                    default_info[key] = value
            else:
                default_info[key] = value
    
    return default_info

def extract_job_info(text: str) -> dict:
    """
    Extract structured job information from job description text using Claude.
    
    Args:
        text: The extracted text from the job description
        
    Returns:
        dict: Structured job information
    """
    if not anthropic_client:
        logger.warning("Anthropic client not available, returning empty job info")
        return post_process_job_info({})
    
    try:
        prompt = f"""Extract job information from the following job description text and return it in the specified JSON schema. 
If a field is missing from the job description, still include it with null, "", or [] as appropriate. Do not skip any fields.

Job Description Text:
{text}

Return the information in this exact JSON schema:
{{
    "job_title": "string or null",
    "location": "string or null",
    "company": "string or null",
    "department": "string or null",
    "employment_type": "string or null",
    "summary": "string or null",
    "responsibilities": ["string"],
    "requirements": ["string"],
    "skills": ["string"],
    "salary": "string or null",
    "benefits": ["string"]
}}

Return only the JSON object, no other text or explanation."""

        response = anthropic_client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )
        
        if not response.content:
            logger.warning("Empty response from Claude for job info extraction")
            return post_process_job_info({})
            
        try:
            # Extract the JSON from the response
            content = response.content[0].text
            # Find the first { and last } to extract the JSON
            start = content.find('{')
            end = content.rfind('}') + 1
            if start == -1 or end == 0:
                logger.warning("No JSON found in Claude response")
                return post_process_job_info({})
                
            json_str = content[start:end]
            extracted_info = json.loads(json_str)
            return post_process_job_info(extracted_info)
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse Claude response as JSON: {e}")
            return post_process_job_info({})
            
    except Exception as e:
        logger.error(f"Error extracting job info with Claude: {e}")
        return post_process_job_info({})

def post_process_job_info(info: dict) -> dict:
    """
    Ensure all expected fields exist in the extracted job information.
    
    Args:
        info: The extracted job information
        
    Returns:
        dict: Post-processed job information with all fields
    """
    default_info = {
        "job_title": None,
        "location": None,
        "company": None,
        "department": None,
        "employment_type": None,
        "summary": None,
        "responsibilities": [],
        "requirements": [],
        "skills": [],
        "salary": None,
        "benefits": []
    }
    
    # Update with any existing values from info
    for key, value in info.items():
        if key in default_info:
            default_info[key] = value
    
    return default_info

def extract_structured_info(text: str, doc_type: str = "job") -> Dict[str, Any]:
    """
    Extract structured information from text using Claude.
    
    Args:
        text (str): The text to extract information from
        doc_type (str): Either "job" or "candidate" to determine extraction type
    
    Returns:
        Dict containing the extracted structured information
    """
    if not anthropic_client:
        logger.warning("Anthropic client not available, returning empty structured info")
        return {} if doc_type == "job" else {}

    try:
        system_prompt = ""
        if doc_type == "job":
            system_prompt = """You are an expert at parsing job descriptions. Extract the following information in JSON format:
            - title: The job title
            - company: Company name
            - location: Job location
            - department: Department or team
            - employment_type: Full-time, part-time, contract, etc.
            - summary: Brief job summary
            - responsibilities: List of key responsibilities
            - requirements: List of requirements/qualifications
            - skills: List of required skills
            - salary: Salary information if available
            - benefits: List of benefits
            
            Return ONLY the JSON object with these fields. Use null for missing fields."""
        else:
            system_prompt = """You are an expert at parsing resumes. Extract the following information in JSON format:
            - name: Candidate's full name
            - email: Email address
            - phone: Phone number
            - location: Location
            - summary: Professional summary
            - education: List of education entries (school, degree, dates)
            - experience: List of work experience entries with the following structure:
                {
                    "job_title": "string",
                    "company": "string",
                    "duration": "string (e.g., '2015-2020', '2 years', '6 months', 'Jan 2015 - Dec 2020')",
                    "responsibilities": ["string"]
                }
            - skills: List of skills
            - languages: List of languages
            - certifications: List of certifications
            
            Important for experience:
            1. Look for sections like 'Work Experience', 'Professional Experience', 'Employment History'
            2. Extract dates in any format (e.g., '2015-2020', '2 years', '6 months', 'Jan 2015 - Dec 2020')
            3. Include all job positions, even if dates are not clear
            4. For responsibilities, extract bullet points or paragraphs describing the role
            5. If no experience section is found, return an empty array []
            6. For each experience entry:
               - Extract the job title and company name
               - Find and format the duration in a consistent way
               - Extract all responsibilities as separate bullet points
               - If duration is unclear, use "Duration not specified"
               - If responsibilities are unclear, use an empty array
            7. IMPORTANT: Only include professional work experience. Do NOT include:
               - Education/tertiary years
               - Internships or student work unless explicitly professional
               - Volunteer work unless it's directly relevant to the profession
               - Part-time work during education unless it's professional experience
            
            Return ONLY the JSON object with these fields. Use null for missing fields."""

        message = anthropic_client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4000,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": text
                }
            ]
        )

        if not message.content:
            logger.warning("Empty response from Claude")
            return {} if doc_type == "job" else {}

        try:
            # Extract the JSON from the response
            content = message.content[0].text
            # Find the first { and last } to extract the JSON
            start = content.find('{')
            end = content.rfind('}') + 1
            if start == -1 or end == 0:
                logger.warning("No JSON found in Claude response")
                return {} if doc_type == "job" else {}
                
            json_str = content[start:end]
            extracted_info = json.loads(json_str)
            
            # Log the raw response for debugging
            logger.info(f"Raw Claude response: {content}")
            logger.info(f"Extracted JSON: {json.dumps(extracted_info, indent=2)}")
            
            # Ensure all required fields exist and have the correct type
            if doc_type == "candidate":
                # Ensure experience is always an array
                if "experience" not in extracted_info:
                    extracted_info["experience"] = []
                elif extracted_info["experience"] is None:
                    extracted_info["experience"] = []
                
                # Ensure each experience entry has the required fields
                for exp in extracted_info["experience"]:
                    if "job_title" not in exp:
                        exp["job_title"] = ""
                    if "company" not in exp:
                        exp["company"] = ""
                    if "duration" not in exp:
                        exp["duration"] = "Duration not specified"
                    if "responsibilities" not in exp:
                        exp["responsibilities"] = []
                    elif exp["responsibilities"] is None:
                        exp["responsibilities"] = []
                
                # Ensure other fields exist
                if "name" not in extracted_info:
                    extracted_info["name"] = None
                if "email" not in extracted_info:
                    extracted_info["email"] = None
                if "phone" not in extracted_info:
                    extracted_info["phone"] = None
                if "location" not in extracted_info:
                    extracted_info["location"] = None
                if "summary" not in extracted_info:
                    extracted_info["summary"] = None
                if "education" not in extracted_info:
                    extracted_info["education"] = []
                if "skills" not in extracted_info:
                    extracted_info["skills"] = []
                if "languages" not in extracted_info:
                    extracted_info["languages"] = []
                if "certifications" not in extracted_info:
                    extracted_info["certifications"] = []
            
            # Log the processed info for debugging
            logger.info(f"Processed info: {json.dumps(extracted_info, indent=2)}")
            
            if doc_type == "job":
                return JobInfo(**extracted_info).model_dump()
            else:
                return CandidateInfo(**extracted_info).model_dump()
                
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse Claude response as JSON: {e}")
            logger.error(f"Response content: {message.content[0].text}")
            return {} if doc_type == "job" else {}
            
    except Exception as e:
        logger.error(f"Error extracting structured info with Claude: {e}")
        return {} if doc_type == "job" else {}

def parse_document(content: bytes, content_type: str, doc_type: str = "job") -> Tuple[str, Dict[str, Any]]:
    """
    Parse a document and return cleaned text and metadata.
    
    Args:
        content (bytes): The document content
        content_type (str): The MIME type of the document
        doc_type (str): Either "job" or "candidate" to determine extraction type
    
    Returns:
        Tuple of (cleaned_text, metadata_dict)
    """
    logger.info(f"Parsing document of type: {content_type}")
    
    try:
        text = ""
        if content_type == "application/pdf":
            # Parse PDF using PyMuPDF
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Parse DOCX using python-docx
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

        if not text.strip():
            raise Exception("No text could be extracted from document")

        # Clean and deduplicate text
        cleaned_text = clean_text(text)
        word_count = len(cleaned_text.split())
        
        # Calculate parse score based on multiple factors
        def calculate_parse_score(text: str, doc_type: str, extracted_info: dict) -> float:
            # Base score starts at 40
            score = 40
            
            # 1. Word count factor (max 10 points)
            word_count = len(text.split())
            if word_count >= 1000:
                score += 10
            elif word_count >= 500:
                score += 7
            elif word_count >= 200:
                score += 5
            elif word_count >= 100:
                score += 3
            
            # 2. Structure factor (max 10 points)
            paragraphs = text.split('\n\n')
            if len(paragraphs) >= 10:
                score += 10
            elif len(paragraphs) >= 5:
                score += 7
            elif len(paragraphs) >= 3:
                score += 5
            
            if doc_type == 'job':
                # 3. Content quality factor for jobs (max 20 points)
                job_terms = ['experience', 'skills', 'requirements', 'responsibilities', 'qualifications']
                term_matches = sum(1 for term in job_terms if term.lower() in text.lower())
                score += min(20, term_matches * 4)
                
                # 4. Information extraction factor for jobs (max 20 points)
                if extracted_info:
                    fields = ['job_title', 'company', 'responsibilities', 'requirements', 'skills']
                    field_score = sum(4 for field in fields if extracted_info.get(field))
                    score += min(20, field_score)
            else:
                # 3. Content quality factor for CVs (max 20 points)
                cv_score = 0
                
                # Check for essential CV sections
                cv_sections = {
                    'contact': ['email', 'phone', 'location'],
                    'experience': ['experience', 'work', 'employment', 'history'],
                    'education': ['education', 'qualification', 'degree', 'university'],
                    'skills': ['skills', 'competencies', 'expertise']
                }
                
                for section, terms in cv_sections.items():
                    if any(term.lower() in text.lower() for term in terms):
                        cv_score += 5
                
                score += min(20, cv_score)
                
                # 4. Information extraction factor for CVs (max 20 points)
                if extracted_info:
                    cv_field_score = 0
                    
                    # Contact information (5 points)
                    contact_fields = ['name', 'email', 'phone']
                    if any(extracted_info.get(field) for field in contact_fields):
                        cv_field_score += 5
                    
                    # Experience quality (8 points)
                    experience = extracted_info.get('experience', [])
                    if experience:
                        exp_score = 0
                        for exp in experience:
                            if exp.get('job_title') and exp.get('company'):
                                exp_score += 2
                            if exp.get('duration'):
                                exp_score += 1
                            if exp.get('responsibilities'):
                                exp_score += 1
                        cv_field_score += min(8, exp_score)
                    
                    # Education (4 points)
                    education = extracted_info.get('education', [])
                    if education and any(edu.get('degree') and edu.get('institution') for edu in education):
                        cv_field_score += 4
                    
                    # Skills (3 points)
                    if extracted_info.get('skills'):
                        cv_field_score += 3
                    
                    score += min(20, cv_field_score)
            
            # Ensure score is between 0 and 100
            return max(0, min(100, score))
        
        # Extract structured information
        extracted_info = extract_structured_info(cleaned_text, doc_type)
        
        # Calculate parse score
        parse_score = calculate_parse_score(cleaned_text, doc_type, extracted_info)
        preview = cleaned_text[:500] + "..." if len(cleaned_text) > 500 else cleaned_text

        metadata = {
            "content_type": content_type,
            "text": cleaned_text,
            "word_count": word_count,
            "parse_score": parse_score,
            "preview": preview,
            "extracted_info": extracted_info,
            "created_at": datetime.utcnow()
        }

        return cleaned_text, metadata

    except Exception as e:
        logger.error(f"Error parsing document: {str(e)}")
        raise
